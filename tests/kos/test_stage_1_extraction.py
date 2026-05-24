"""
Tests for Stage 1 — document extraction.

Covers PDF native, DOCX with tables, and image (OCR path).
"""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest

from services.ingestion.stage_1_extraction import extract_structure


class TestPdfNative:
    def test_extracts_text_from_native_pdf(self, sample_pdf):
        doc = extract_structure(str(sample_pdf), "pdf")
        assert "Hello KOS" in doc.raw_text
        assert doc.is_scanned is False
        assert len(doc.sections) >= 1

    def test_language_detected(self, sample_pdf):
        doc = extract_structure(str(sample_pdf), "pdf")
        assert doc.language in {"en", "und"}  # English or undetermined

    def test_markdown_not_empty(self, sample_pdf):
        doc = extract_structure(str(sample_pdf), "pdf")
        assert len(doc.markdown) > 0

    def test_source_tool_set(self, sample_pdf):
        doc = extract_structure(str(sample_pdf), "pdf")
        assert doc.source_tool != ""


class TestPdfScanned:
    @pytest.fixture()
    def scanned_pdf(self, tmp_path) -> Path:
        """Create a PDF that looks like a scan (image embedded, no text layer)."""
        from PIL import Image, ImageDraw  # noqa: PLC0415
        import io  # noqa: PLC0415

        img = Image.new("RGB", (800, 1100), color="white")
        draw = ImageDraw.Draw(img)
        draw.text((100, 200), "Invoice total: $1,234.00", fill="black")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)

        doc = fitz.open()
        page = doc.new_page(width=800, height=1100)
        page.insert_image(page.rect, stream=buf.read())
        path = tmp_path / "scanned.pdf"
        doc.save(str(path))
        doc.close()
        return path

    @pytest.mark.skipif(
        not _tesseract_available(), reason="Tesseract OCR not installed"
    )
    def test_scanned_pdf_ocr(self, scanned_pdf):
        doc = extract_structure(str(scanned_pdf), "pdf")
        assert doc.is_scanned is True
        # OCR may not be perfect; just check something was extracted
        assert len(doc.raw_text) > 0


class TestDocx:
    @pytest.fixture()
    def sample_docx(self, tmp_path) -> Path:
        pytest.importorskip("docx", reason="python-docx not installed")
        from docx import Document  # noqa: PLC0415
        from docx.oxml.ns import qn  # noqa: PLC0415

        doc = Document()
        doc.add_heading("Contract Terms", level=1)
        doc.add_paragraph("This agreement is effective as of January 1, 2025.")
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Party"
        tbl.cell(0, 1).text = "Role"
        tbl.cell(1, 0).text = "Acme Corp"
        tbl.cell(1, 1).text = "Vendor"
        path = tmp_path / "contract.docx"
        doc.save(str(path))
        return path

    def test_extracts_docx_text(self, sample_docx):
        doc = extract_structure(str(sample_docx), "docx")
        assert "Contract Terms" in doc.markdown or "Contract Terms" in doc.raw_text

    def test_docx_has_sections(self, sample_docx):
        doc = extract_structure(str(sample_docx), "docx")
        assert len(doc.sections) >= 1


class TestXlsx:
    @pytest.fixture()
    def sample_xlsx(self, tmp_path) -> Path:
        pytest.importorskip("openpyxl")
        import openpyxl  # noqa: PLC0415
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sales"
        ws.append(["Month", "Revenue"])
        ws.append(["January", 50000])
        ws.append(["February", 62000])
        path = tmp_path / "sales.xlsx"
        wb.save(str(path))
        return path

    def test_xlsx_extracted_as_tables(self, sample_xlsx):
        doc = extract_structure(str(sample_xlsx), "xlsx")
        assert len(doc.tables) >= 1
        assert doc.tables[0].title == "Sales"
        assert doc.tables[0].data[0] == ["Month", "Revenue"]

    def test_xlsx_raw_text_contains_values(self, sample_xlsx):
        doc = extract_structure(str(sample_xlsx), "xlsx")
        assert "50000" in doc.raw_text or "January" in doc.raw_text


class TestUnsupportedType:
    def test_unsupported_type_raises(self, sample_pdf):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_structure(str(sample_pdf), "mp4")


def _tesseract_available() -> bool:
    import shutil  # noqa: PLC0415
    return shutil.which("tesseract") is not None
